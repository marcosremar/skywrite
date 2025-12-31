#!/usr/bin/env python3
"""
LaTeX to Word Converter - Enhanced Version with Comprehensive Formatting

This script converts a LaTeX thesis document to a Word (.docx) format
while preserving formatting as closely as possible to the original.

Usage: python Latex2word.py

Requirements:
    pip install python-docx pylatexenc regex
"""

import os
import re
import sys
from pathlib import Path
from typing import Dict, List, Tuple, Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from docx.enum.dml import MSO_THEME_COLOR_INDEX
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

try:
    from pylatexenc.latex2text import LatexNodes2Text
except ImportError:
    print("Error: pylatexenc not installed. Run: pip install pylatexenc")
    sys.exit(1)

try:
    import regex as re_unicode
except ImportError:
    print("Warning: regex not installed. Using standard re")
    re_unicode = re


class LatexToWordConverter:
    """Converts LaTeX thesis to Word document with comprehensive formatting."""
    
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
        self.latex_converter = LatexNodes2Text()
        self.bibliography = {}
        self.citations = set()
        self.missing_citations = set()  # Track citations not found in bibliography
        self.reference_counter = 1
        self.section_counters = {'chapter': 0, 'section': 0, 'subsection': 0, 'subsubsection': 0}
        
    def setup_styles(self):
        """Set up Word styles to match classic thesis formatting."""
        styles = self.doc.styles
        
        # Configure Normal style (justified, line spacing 1.5, Times New Roman 12pt)
        normal_style = styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(12)
        normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal_style.paragraph_format.line_spacing = 1.5
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.first_line_indent = Pt(14)  # Paragraph indent
        
        # Title style (centered, very large, bold)
        if 'Title' not in [s.name for s in styles]:
            title_style = styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
        else:
            title_style = styles['Title']
        title_style.font.name = 'Times New Roman'
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_before = Pt(72)
        title_style.paragraph_format.space_after = Pt(72)
        title_style.paragraph_format.first_line_indent = Pt(0)
        
        # Chapter style (centered, large, bold)
        if 'Chapter' not in [s.name for s in styles]:
            chapter_style = styles.add_style('Chapter', WD_STYLE_TYPE.PARAGRAPH)
        else:
            chapter_style = styles['Chapter']
        chapter_style.font.name = 'Times New Roman'
        chapter_style.font.size = Pt(18)
        chapter_style.font.bold = True
        chapter_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chapter_style.paragraph_format.space_before = Pt(36)
        chapter_style.paragraph_format.space_after = Pt(24)
        chapter_style.paragraph_format.first_line_indent = Pt(0)
        
        # Configure heading styles with proper hierarchy
        heading_configs = [
            (1, 16, True),   # Main sections
            (2, 14, True),   # Subsections  
            (3, 13, True),   # Sub-subsections
            (4, 12, True),   # Paragraphs
        ]
        
        for level, size, bold in heading_configs:
            style_name = f'Heading {level}'
            if style_name in [s.name for s in styles]:
                style = styles[style_name]
            else:
                style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            
            style.font.name = 'Times New Roman'
            style.font.size = Pt(size)
            style.font.bold = bold
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            style.paragraph_format.space_before = Pt(18 - level * 2)
            style.paragraph_format.space_after = Pt(12 - level * 2)
            style.paragraph_format.first_line_indent = Pt(0)
            style.paragraph_format.keep_with_next = True
        
        # Quote style (indented, smaller font)
        if 'Quote' not in [s.name for s in styles]:
            quote_style = styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
        else:
            quote_style = styles['Quote']
        quote_style.font.name = 'Times New Roman'
        quote_style.font.size = Pt(11)
        quote_style.font.italic = True
        quote_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        quote_style.paragraph_format.left_indent = Inches(0.5)
        quote_style.paragraph_format.right_indent = Inches(0.5)
        quote_style.paragraph_format.space_before = Pt(6)
        quote_style.paragraph_format.space_after = Pt(6)
        quote_style.paragraph_format.first_line_indent = Pt(0)
        
        # Bibliography style with hanging indent
        if 'Bibliography' not in [s.name for s in styles]:
            bib_style = styles.add_style('Bibliography', WD_STYLE_TYPE.PARAGRAPH)
        else:
            bib_style = styles['Bibliography']
        bib_style.font.name = 'Times New Roman'
        bib_style.font.size = Pt(11)
        bib_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        bib_style.paragraph_format.left_indent = Inches(0.5)
        bib_style.paragraph_format.first_line_indent = Inches(-0.5)  # Hanging indent
        bib_style.paragraph_format.space_after = Pt(6)
        
        # Definition style
        if 'Definition' not in [s.name for s in styles]:
            def_style = styles.add_style('Definition', WD_STYLE_TYPE.PARAGRAPH)
        else:
            def_style = styles['Definition']
        def_style.font.name = 'Times New Roman'
        def_style.font.size = Pt(12)
        def_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        def_style.paragraph_format.left_indent = Inches(0.25)
        def_style.paragraph_format.space_before = Pt(6)
        def_style.paragraph_format.space_after = Pt(6)
        def_style.paragraph_format.first_line_indent = Pt(0)
    
    def process_bibliography_files(self):
        """Process bibliography files and extract references."""
        # Use only the Zotero bibliography file as requested
        bib_file = 'bibliography/references-zotero.bib'
        
        if os.path.exists(bib_file):
            print(f"Processing bibliography: {bib_file}")
            content = self.read_file_safely(Path(bib_file))
            self.parse_bibliography_content(content)
        else:
            print(f"Warning: Bibliography file not found: {bib_file}")
    
    def parse_bibliography_content(self, content: str):
        """Parse bibliography content with optimized approach."""
        # Simpler, faster regex for BibTeX entries
        entry_pattern = r'@(\w+)\s*\{\s*([^,\s]+)\s*,([^@]*?)(?=\n@|\Z)'
        
        entries_found = 0
        for match in re.finditer(entry_pattern, content, re.DOTALL):
            entries_found += 1
            if entries_found > 500:  # Limit for performance
                print(f"Limiting bibliography processing to first 500 entries for performance")
                break
                
            entry_type = match.group(1)
            entry_key = match.group(2)
            entry_fields = match.group(3)
            
            # Extract key fields quickly
            fields = {}
            simple_fields = ['author', 'title', 'year', 'date', 'journal', 'journaltitle', 
                           'publisher', 'booktitle', 'volume', 'pages']
            
            for field_name in simple_fields:
                pattern = rf'{field_name}\s*=\s*\{{([^}}]*)\}}'
                match_field = re.search(pattern, entry_fields, re.IGNORECASE)
                if match_field:
                    fields[field_name.lower()] = match_field.group(1)
            
            # Format reference
            formatted_ref = self.format_reference(entry_type, fields)
            self.bibliography[entry_key] = formatted_ref
    
    def format_reference(self, entry_type: str, fields: dict) -> str:
        """Format a bibliography reference in academic style."""
        author = fields.get('author', 'Unknown Author')
        title = fields.get('title', 'Untitled')
        year = fields.get('date', fields.get('year', 'n.d.'))
        
        # Clean up fields
        author = re.sub(r'\{|\}', '', author).strip()
        title = re.sub(r'\{|\}', '', title).strip()
        year = re.sub(r'\{|\}', '', year).strip()
        
        # Extract year from date if needed
        if '-' in year:
            year = year.split('-')[0]
        
        # Simplify author names (surname, initials)
        if ' and ' in author:
            authors = author.split(' and ')
            author = authors[0]  # Use first author for citation
        
        if entry_type.lower() == 'article':
            journal = fields.get('journaltitle', fields.get('journal', ''))
            journal = re.sub(r'\{|\}', '', journal).strip()
            volume = fields.get('volume', '')
            pages = fields.get('pages', '')
            
            ref = f"{author} ({year}). {title}."
            if journal:
                ref += f" *{journal}*"
            if volume:
                ref += f", {volume}"
            if pages:
                ref += f", {pages}"
            return ref + "."
        
        elif entry_type.lower() == 'book':
            publisher = fields.get('publisher', '')
            publisher = re.sub(r'\{|\}', '', publisher).strip()
            location = fields.get('location', fields.get('address', ''))
            location = re.sub(r'\{|\}', '', location).strip()
            
            ref = f"{author} ({year}). *{title}*."
            if publisher:
                if location:
                    ref += f" {location}: {publisher}."
                else:
                    ref += f" {publisher}."
            return ref
        
        elif entry_type.lower() in ['inproceedings', 'incollection']:
            booktitle = fields.get('booktitle', '')
            booktitle = re.sub(r'\{|\}', '', booktitle).strip()
            editor = fields.get('editor', '')
            editor = re.sub(r'\{|\}', '', editor).strip()
            
            ref = f"{author} ({year}). {title}."
            if booktitle:
                ref += f" In"
                if editor:
                    ref += f" {editor} (Ed.),"
                ref += f" *{booktitle}*."
            return ref
        
        else:
            return f"{author} ({year}). {title}."
    
    def clean_latex_text(self, text: str) -> str:
        """Clean LaTeX commands and convert to plain text with enhanced processing."""
        if not text:
            return ""
        
        # Remove comments first
        text = re.sub(r'%.*$', '', text, flags=re.MULTILINE)
        
        # Handle citations first (before general cleaning)
        text = self.process_citations(text)
        
        # Convert special LaTeX formatting
        replacements = [
            # Text formatting
            (r'\\textbf\{([^}]+)\}', r'**\1**'),  # Bold - will be processed later
            (r'\\textit\{([^}]+)\}', r'*\1*'),    # Italic - will be processed later  
            (r'\\emph\{([^}]+)\}', r'*\1*'),      # Emphasis - will be processed later
            (r'\\textsc\{([^}]+)\}', r'\1'),      # Small caps to normal
            
            # Special characters
            (r'\\&', '&'),
            (r'\\\$', '$'),
            (r'\\%', '%'),
            (r'\\_', '_'),
            (r'\\#', '#'),
            (r'\\textasciitilde', '~'),
            (r'\\textbackslash', r'\\'),
            
            # Spacing and breaks
            (r'\\\\', '\n'),
            (r'\\newline', '\n'),
            (r'\\newpage', '\n\n'),
            (r'\\clearpage', '\n\n'),
            (r'\\par', '\n\n'),
            (r'\\noindent', ''),
            (r'\\indent', ''),
            
            # Math and symbols
            (r'\\ldots', '...'),
            (r'\\dots', '...'),
            (r'\\LaTeX', 'LaTeX'),
            (r'\\TeX', 'TeX'),
            
            # Spaces
            (r'\\,', ' '),
            (r'\\;', ' '),
            (r'\\quad', '    '),
            (r'\\qquad', '        '),
            (r'\\xspace', ''),
            
            # Remove problematic commands
            (r'\\@', ''),
            (r'\\label\{[^}]*\}', ''),
            (r'\\ref\{[^}]*\}', '[REF]'),
            (r'\\pageref\{[^}]*\}', '[PAGE]'),
            (r'\\cite\{[^}]*\}', '[CITE]'),  # Fallback for missed citations
        ]
        
        for pattern, replacement in replacements:
            text = re.sub(pattern, replacement, text)
        
        # Use pylatexenc for remaining conversions
        try:
            text = self.latex_converter.latex_to_text(text)
        except Exception as e:
            print(f"Warning: Could not fully convert LaTeX text: {e}")
        
        # Clean up multiple spaces and newlines
        text = re.sub(r' +', ' ', text)
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        return text.strip()
    
    def process_citations(self, text: str) -> str:
        """Process citations and replace with formatted references."""
        citation_patterns = [
            r'\\autocite\{([^}]+)\}',
            r'\\cite\{([^}]+)\}',
            r'\\textcite\{([^}]+)\}',
            r'\\parencite\{([^}]+)\}',
            r'\\citep\{([^}]+)\}',
            r'\\citet\{([^}]+)\}'
        ]
        
        def replace_citation(match):
            cite_keys = match.group(1).split(',')
            formatted_cites = []
            
            for key in cite_keys:
                key = key.strip()
                self.citations.add(key)
                
                if key in self.bibliography:
                    ref = self.bibliography[key]
                    # Extract author and year from bibliography
                    author_year_match = re.search(r'^([^(]+)\s*\((\d{4})\)', ref)
                    if author_year_match:
                        author = author_year_match.group(1).strip()
                        year = author_year_match.group(2)
                        
                        # Simplify author name (last name only)
                        if ',' in author:
                            author = author.split(',')[0].strip()
                        elif ' ' in author:
                            parts = author.split()
                            author = parts[-1]  # Last name
                        
                        formatted_cites.append(f"{author}, {year}")
                    else:
                        formatted_cites.append(key)
                else:
                    # Mark as missing citation for red formatting
                    self.missing_citations.add(key)
                    formatted_cites.append(f"***MISSING_REF:{key}***")
            
            if len(formatted_cites) == 1:
                return f"({formatted_cites[0]})"
            else:
                return f"({'; '.join(formatted_cites)})"
        
        for pattern in citation_patterns:
            text = re.sub(pattern, replace_citation, text)
        
        return text
    
    def add_formatted_paragraph(self, text: str, style_name: str = 'Normal'):
        """Add paragraph with advanced formatting to document."""
        if not text.strip():
            return
        
        paragraph = self.doc.add_paragraph(style=style_name)
        
        # Process text for inline formatting
        self.process_inline_formatting(paragraph, text)
    
    def process_inline_formatting(self, paragraph, text: str):
        """Process inline formatting like bold, italic, and missing references."""
        # Clean the text first
        clean_text = self.clean_latex_text(text)
        
        # Split text by formatting markers and missing references
        parts = []
        current_text = clean_text
        
        # Process missing references first (highest priority)
        missing_ref_pattern = r'\*\*\*MISSING_REF:([^*]+)\*\*\*'
        while re.search(missing_ref_pattern, current_text):
            match = re.search(missing_ref_pattern, current_text)
            if match:
                # Add text before match
                before = current_text[:match.start()]
                if before:
                    parts.append((before, False, False, False))  # text, bold, italic, is_missing_ref
                
                # Add just the reference key in red (no extra text)
                parts.append((match.group(1), False, False, True))
                
                # Continue with rest
                current_text = current_text[match.end():]
            else:
                break
        
        # Process bold (**text**)
        bold_pattern = r'\*\*([^*]+)\*\*'
        temp_parts = []
        for part_text, is_bold, is_italic, is_missing in parts + [(current_text, False, False, False)]:
            if not is_bold and not is_italic and not is_missing:  # Only process if not already formatted
                while re.search(bold_pattern, part_text):
                    match = re.search(bold_pattern, part_text)
                    if match:
                        # Add text before match
                        before = part_text[:match.start()]
                        if before:
                            temp_parts.append((before, False, False, False))
                        
                        # Add bold text
                        temp_parts.append((match.group(1), True, False, False))
                        
                        # Continue with rest
                        part_text = part_text[match.end():]
                    else:
                        break
                
                if part_text:
                    temp_parts.append((part_text, is_bold, is_italic, is_missing))
            else:
                temp_parts.append((part_text, is_bold, is_italic, is_missing))
        
        parts = temp_parts
        
        # Process italic (*text*)
        italic_pattern = r'\*([^*]+)\*'
        temp_parts = []
        for part_text, is_bold, is_italic, is_missing in parts:
            if not is_bold and not is_italic and not is_missing:  # Only process if not already formatted
                while re.search(italic_pattern, part_text):
                    match = re.search(italic_pattern, part_text)
                    if match:
                        # Add text before match
                        before = part_text[:match.start()]
                        if before:
                            temp_parts.append((before, is_bold, False, is_missing))
                        
                        # Add italic text
                        temp_parts.append((match.group(1), is_bold, True, is_missing))
                        
                        # Continue with rest
                        part_text = part_text[match.end():]
                    else:
                        break
                
                if part_text:
                    temp_parts.append((part_text, is_bold, is_italic, is_missing))
            else:
                temp_parts.append((part_text, is_bold, is_italic, is_missing))
        
        parts = temp_parts
        
        # Add runs to paragraph
        if not parts and clean_text:
            parts = [(clean_text, False, False, False)]
        
        for part_text, is_bold, is_italic, is_missing_ref in parts:
            if part_text.strip():
                # Clean text to avoid XML errors
                clean_part_text = self.clean_text_for_word(part_text)
                if clean_part_text:
                    run = paragraph.add_run(clean_part_text)
                    run.bold = is_bold
                    run.italic = is_italic
                    
                    # Color missing references red
                    if is_missing_ref:
                        run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
    
    def clean_text_for_word(self, text: str) -> str:
        """Clean text to avoid XML compatibility issues while preserving Portuguese characters."""
        if not text:
            return ""
        
        # Define control characters to remove (but keep Portuguese special chars)
        control_chars = [
            '\x00', '\x01', '\x02', '\x03', '\x04', '\x05', '\x06', '\x07',
            '\x08', '\x0b', '\x0c', '\x0e', '\x0f', '\x10', '\x11', '\x12',
            '\x13', '\x14', '\x15', '\x16', '\x17', '\x18', '\x19', '\x1a',
            '\x1b', '\x1c', '\x1d', '\x1e', '\x1f'
        ]
        
        # Remove only problematic control characters, preserve Portuguese chars
        cleaned = text
        for char in control_chars:
            cleaned = cleaned.replace(char, '')
        
        # Also remove any character that's not printable but preserve Portuguese accents
        import unicodedata
        result = ""
        for char in cleaned:
            # Keep all normal characters, Portuguese accents, and common symbols
            if (char.isprintable() or 
                char in '\n\t\r' or  # Keep important whitespace
                unicodedata.category(char).startswith('L') or  # All letters (including accented)
                unicodedata.category(char).startswith('M') or  # Marks (accents)
                unicodedata.category(char).startswith('N') or  # Numbers
                unicodedata.category(char).startswith('P') or  # Punctuation
                unicodedata.category(char).startswith('S') or  # Symbols
                unicodedata.category(char).startswith('Z')):   # Separators/spaces
                result += char
        
        return result.strip()
    
    def process_table(self, table_content: str):
        """Process LaTeX table and convert to Word table."""
        print("Processing table...")
        
        # Extract table rows and clean up
        rows = []
        lines = table_content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line or line.startswith('%'):
                continue
            
            # Skip LaTeX table commands
            if any(cmd in line for cmd in ['\\begin{', '\\end{', '\\hline', '\\toprule', '\\midrule', '\\bottomrule']):
                continue
            
            # Process table row (split by &)
            if '&' in line:
                # Clean up the row
                row = line.replace('\\\\', '').strip()
                cells = [cell.strip() for cell in row.split('&')]
                
                # Clean each cell
                clean_cells = []
                for cell in cells:
                    clean_cell = self.clean_latex_text(cell)
                    clean_cells.append(clean_cell)
                
                if clean_cells:
                    rows.append(clean_cells)
        
        # Create Word table if we have rows
        if rows:
            # Determine table dimensions
            max_cols = max(len(row) for row in rows) if rows else 0
            num_rows = len(rows)
            
            if max_cols > 0 and num_rows > 0:
                # Create table
                table = self.doc.add_table(rows=num_rows, cols=max_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = 'Table Grid'
                
                # Fill table with data
                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < max_cols:
                            cell = table.cell(row_idx, col_idx)
                            cell.text = cell_data
                            
                            # Format first row as header
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                
                print(f"Created table with {num_rows} rows and {max_cols} columns")
                return True
        
        return False
    
    def extract_and_process_tables(self, content: str) -> str:
        """Extract LaTeX tables and process them, returning content without tables."""
        # Find table environments
        table_patterns = [
            r'\\begin\{table\}.*?\\end\{table\}',
            r'\\begin\{tabular\}.*?\\end\{tabular\}',
            r'\\begin\{longtable\}.*?\\end\{longtable\}',
            r'\\begin\{array\}.*?\\end\{array\}'
        ]
        
        for pattern in table_patterns:
            matches = list(re.finditer(pattern, content, re.DOTALL))
            for match in matches:
                table_content = match.group(0)
                
                # Try to process the table
                if self.process_table(table_content):
                    # Replace table with placeholder
                    content = content.replace(table_content, '\n[TABLE PROCESSED]\n')
                else:
                    # If table processing failed, try to extract simple data
                    self.process_simple_table_data(table_content)
                    content = content.replace(table_content, '\n[TABLE DATA EXTRACTED]\n')
        
        return content
    
    def process_simple_table_data(self, table_content: str):
        """Process simple table data that might not be in formal table environment."""
        # Look for lines that might be table data (contain & or multiple columns)
        lines = table_content.split('\n')
        table_lines = []
        
        for line in lines:
            line = line.strip()
            # Check if line looks like table data
            if ('&' in line or 
                (len(line.split()) >= 3 and not line.startswith('\\'))):
                table_lines.append(line)
        
        # If we found potential table data, create a simple table
        if len(table_lines) >= 2:  # At least header and one data row
            print(f"Processing simple table data with {len(table_lines)} rows")
            
            # Try to create table from this data
            rows = []
            for line in table_lines:
                if '&' in line:
                    # LaTeX table format
                    cells = [cell.strip() for cell in line.split('&')]
                    cells = [self.clean_latex_text(cell.replace('\\\\', '')) for cell in cells]
                else:
                    # Space-separated data
                    cells = line.split()
                    cells = [self.clean_latex_text(cell) for cell in cells if cell]
                
                if cells:
                    rows.append(cells)
            
            # Create Word table
            if rows:
                max_cols = max(len(row) for row in rows)
                table = self.doc.add_table(rows=len(rows), cols=max_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = 'Table Grid'
                
                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < max_cols:
                            cell = table.cell(row_idx, col_idx)
                            cell.text = cell_data
                            
                            # Format first row as header
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                
                print(f"Created simple table with {len(rows)} rows and {max_cols} columns")
    
    def read_file_safely(self, filepath: Path) -> str:
        """Read file with proper encoding handling."""
        encodings = ['utf-8', 'latin1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
            except FileNotFoundError:
                print(f"Warning: File not found: {filepath}")
                return ""
        
        print(f"Warning: Could not read file {filepath} with any encoding")
        return ""
    
    def process_input_command(self, content: str) -> str:
        """Process \\input{} commands and include file contents."""
        content = self.process_def_commands(content)
        
        input_pattern = r'\\input\{([^}]+)\}'
        
        def replace_input(match):
            filename = match.group(1)
            if not filename.endswith('.tex'):
                filename += '.tex'
            
            filepath = Path(filename)
            if filepath.exists():
                file_content = self.read_file_safely(filepath)
                return self.process_input_command(file_content)
            else:
                print(f"Warning: Input file not found: {filename}")
                return ""
        
        return re.sub(input_pattern, replace_input, content)
    
    def process_def_commands(self, content: str) -> str:
        """Process \\def commands and replace variables."""
        def_pattern = r'\\def\\([^{]+)\{([^}]+)\}'
        defs = {}
        
        for match in re.finditer(def_pattern, content):
            var_name = match.group(1)
            var_value = match.group(2)
            defs[var_name] = var_value
        
        for var_name, var_value in defs.items():
            content = content.replace(f'\\{var_name}', var_value)
        
        return content
    
    def determine_heading_level(self, section_type: str) -> int:
        """Determine the appropriate heading level based on section type."""
        level_map = {
            'chapter': 0,      # Special case - use Chapter style
            'section': 1,      # Heading 1
            'subsection': 2,   # Heading 2
            'subsubsection': 3, # Heading 3
            'paragraph': 4,    # Heading 4
        }
        return level_map.get(section_type, 1)
    
    def process_latex_structure(self, content: str):
        """Process LaTeX structure with improved section handling."""
        # Remove LaTeX commands first, but keep table processing inline
        content = self.remove_latex_commands(content)
        
        # Split into paragraphs but preserve structure
        paragraphs = re.split(r'\n\s*\n', content)
        
        for para in paragraphs:
            para = para.strip()
            if not para or para.startswith('%'):
                continue
            
            # Check for chapter
            chapter_match = re.search(r'\\chapter\{([^}]+)\}', para)
            if chapter_match:
                self.section_counters['chapter'] += 1
                self.section_counters['section'] = 0  # Reset section counter
                
                title = chapter_match.group(1)
                chapter_title = f"Capítulo {self.section_counters['chapter']}: {title}"
                self.add_formatted_paragraph(chapter_title, 'Chapter')
                
                # Add remaining text if any
                remaining = re.sub(r'\\chapter\{[^}]+\}', '', para).strip()
                if remaining:
                    self.add_formatted_paragraph(remaining)
                continue
            
            # Check for sections with numbering
            section_match = re.search(r'\\((?:sub)*)section\{([^}]+)\}', para)
            if section_match:
                section_type = section_match.group(1) + 'section'
                title = section_match.group(2)
                
                # Update counters
                if section_type == 'section':
                    self.section_counters['section'] += 1
                    self.section_counters['subsection'] = 0
                    section_num = f"{self.section_counters['chapter']}.{self.section_counters['section']}"
                elif section_type == 'subsection':
                    self.section_counters['subsection'] += 1
                    section_num = f"{self.section_counters['chapter']}.{self.section_counters['section']}.{self.section_counters['subsection']}"
                else:
                    section_num = ""
                
                # Format title with numbering
                if section_num:
                    formatted_title = f"{section_num} {title}"
                else:
                    formatted_title = title
                
                level = self.determine_heading_level(section_type)
                self.add_formatted_paragraph(formatted_title, f'Heading {level}')
                
                # Add remaining text
                remaining = re.sub(r'\\(?:sub)*section\{[^}]+\}', '', para).strip()
                if remaining:
                    self.add_formatted_paragraph(remaining)
                continue
            
            # Check for paragraph command
            para_match = re.search(r'\\paragraph\{([^}]+)\}', para)
            if para_match:
                title = para_match.group(1)
                self.add_formatted_paragraph(title, 'Heading 4')
                
                remaining = re.sub(r'\\paragraph\{[^}]+\}', '', para).strip()
                if remaining:
                    self.add_formatted_paragraph(remaining)
                continue
            
            # Skip LaTeX commands
            skip_patterns = [
                r'\\documentclass', r'\\usepackage', r'\\newcommand', r'\\renewcommand',
                r'\\def', r'\\let', r'\\begin\{document\}', r'\\end\{document\}',
                r'\\input\{', r'\\include', r'\\frenchspacing', r'\\raggedbottom',
                r'\\selectlanguage', r'\\pagenumbering', r'\\pagestyle',
                r'\\cleardoublepage', r'\\thispagestyle', r'\\togglecommentblocks',
                r'\\pagecolor', r'\\nopagecolor', r'\\addbibresource', r'\\appendix',
                r'\\bookmarksetup'
            ]
            
            if any(re.search(pattern, para) for pattern in skip_patterns):
                continue
            
            # Check for table content in this paragraph
            if self.contains_table(para):
                print(f"Table detected in paragraph: {para[:100]}...")
                self.process_paragraph_with_table(para)
            # Regular paragraph - check for special content
            elif self.is_definition_paragraph(para):
                self.add_formatted_paragraph(para, 'Definition')
            elif self.is_list_paragraph(para):
                self.process_list_paragraph(para)
            elif self.is_informal_table(para):
                self.process_informal_table(para)
            else:
                self.add_formatted_paragraph(para, 'Normal')
    
    def is_definition_paragraph(self, text: str) -> bool:
        """Check if paragraph contains definitions."""
        definition_indicators = [
            r'define-se',
            r'definimos',
            r'conceito de',
            r'entende-se por',
            r'considera-se',
            r'compreende-se'
        ]
        return any(re.search(pattern, text, re.IGNORECASE) for pattern in definition_indicators)
    
    def is_list_paragraph(self, text: str) -> bool:
        """Check if paragraph should be formatted as a list."""
        list_indicators = [
            r'são elas:',
            r'incluem:',
            r'destacam-se:',
            r'principais são:',
            r'podem ser:',
            r'dividem-se em:'
        ]
        return any(re.search(pattern, text, re.IGNORECASE) for pattern in list_indicators)
    
    def process_list_paragraph(self, text: str):
        """Process paragraph that should be formatted as a list."""
        # For now, just add as normal paragraph
        # In the future, could implement actual list formatting
        self.add_formatted_paragraph(text, 'Normal')
    
    def is_informal_table(self, text: str) -> bool:
        """Check if paragraph contains informal table data."""
        # Look for patterns that suggest tabular data
        lines = text.split('\n')
        potential_table_lines = 0
        
        for line in lines:
            line = line.strip()
            words = line.split()
            if (len(words) >= 3 and 
                not line.startswith('\\') and
                not line.lower().startswith(('este', 'esta', 'o', 'a', 'um', 'uma', 'como', 'para'))):
                
                # Check if it looks like structured data
                if any(keyword in line.lower() for keyword in ['capítulo', 'conteúdo', 'questões', 'principal']):
                    potential_table_lines += 1
        
        return potential_table_lines >= 2
    
    def process_informal_table(self, text: str):
        """Process informal table data found in text."""
        print("Processing informal table data...")
        
        lines = text.split('\n')
        table_rows = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            words = line.split()
            
            # Look for lines that seem to be table data
            if (len(words) >= 2 and 
                not line.startswith('\\') and
                any(keyword in line.lower() for keyword in ['capítulo', 'conteúdo', 'questões', 'principal', 'metodologia', 'resultados'])):
                
                # Try to identify column structure
                clean_line = self.clean_latex_text(line)
                if clean_line:
                    # Split by common separators or keep as single cells
                    if ':' in clean_line:
                        cells = [cell.strip() for cell in clean_line.split(':')]
                    elif '–' in clean_line or '—' in clean_line:
                        cells = [cell.strip() for cell in re.split(r'[–—]', clean_line)]
                    else:
                        # Try to split intelligently by space
                        words = clean_line.split()
                        if len(words) <= 3:
                            cells = words
                        else:
                            # Group words into logical cells
                            cells = [' '.join(words[:2]), ' '.join(words[2:])]
                    
                    if len(cells) >= 2:
                        table_rows.append(cells)
        
        # Create table if we found structured data
        if len(table_rows) >= 2:
            max_cols = max(len(row) for row in table_rows)
            
            # Add header if it seems to be missing
            if not any('capítulo' in row[0].lower() for row in table_rows[:1]):
                header = ['Capítulo', 'Conteúdo']
                if max_cols > 2:
                    header.extend([f'Coluna {i}' for i in range(3, max_cols + 1)])
                table_rows.insert(0, header)
            
            # Create Word table
            table = self.doc.add_table(rows=len(table_rows), cols=max_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Table Grid'
            
            for row_idx, row_data in enumerate(table_rows):
                for col_idx, cell_data in enumerate(row_data):
                    if col_idx < max_cols and col_idx < len(row_data):
                        cell = table.cell(row_idx, col_idx)
                        cell.text = str(cell_data)
                        
                        # Format first row as header
                        if row_idx == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
            
            print(f"Created informal table with {len(table_rows)} rows and {max_cols} columns")
        else:
            # If we couldn't create a table, just add as normal text
            self.add_formatted_paragraph(text, 'Normal')
    
    def contains_table(self, text: str) -> bool:
        """Check if paragraph contains table environment."""
        table_patterns = [
            r'\\begin\{table\}',
            r'\\begin\{tabular\}',
            r'\\begin\{longtable\}',
            r'\\begin\{array\}',
            # Also detect inline table-like content
            r'&.*&.*\\\\',  # Multiple & with line ending
        ]
        return any(re.search(pattern, text) for pattern in table_patterns)
    
    def process_paragraph_with_table(self, text: str):
        """Process a paragraph that contains table environment."""
        # Extract table environments and process them
        table_patterns = [
            r'\\begin\{table\}.*?\\end\{table\}',
            r'\\begin\{tabular\}.*?\\end\{tabular\}',
            r'\\begin\{longtable\}.*?\\end\{longtable\}',
            r'\\begin\{array\}.*?\\end\{array\}'
        ]
        
        remaining_text = text
        
        for pattern in table_patterns:
            matches = list(re.finditer(pattern, remaining_text, re.DOTALL))
            for match in matches:
                table_content = match.group(0)
                
                # Add any text before the table
                before_table = remaining_text[:match.start()].strip()
                if before_table:
                    self.add_formatted_paragraph(before_table, 'Normal')
                
                # Process the table
                if self.process_table(table_content):
                    print("Table processed successfully in place")
                else:
                    # If table processing failed, try simple table data
                    self.process_simple_table_data(table_content)
                
                # Continue with text after the table
                remaining_text = remaining_text[match.end():].strip()
        
        # Add any remaining text after all tables
        if remaining_text:
            self.add_formatted_paragraph(remaining_text, 'Normal')
    
    def remove_latex_commands(self, content: str) -> str:
        """Remove LaTeX commands that don't translate to Word."""
        patterns_to_remove = [
            r'\\dictum\[[^\]]*\]\{[^}]*\}',
            r'\\vskip\s+\d+\w*',
            r'\\RequirePackage\{[^}]*\}',
            r'\\frenchspacing',
            r'\\raggedbottom',
            r'\\begin\{[^}]*\}',
            r'\\end\{[^}]*\}',
        ]
        
        for pattern in patterns_to_remove:
            content = re.sub(pattern, '', content)
        
        return content
    
    def add_bibliography(self):
        """Add properly formatted bibliography section."""
        if not self.citations:
            return
        
        # Add bibliography heading
        self.doc.add_page_break()
        self.add_formatted_paragraph("Referências", 'Chapter')
        
        # Sort citations alphabetically
        used_refs = {}
        for cite_key in sorted(self.citations):
            if cite_key in self.bibliography:
                used_refs[cite_key] = self.bibliography[cite_key]
        
        # Add references with proper formatting
        for ref_key, ref_text in used_refs.items():
            paragraph = self.doc.add_paragraph(style='Bibliography')
            
            # Process italic formatting in references
            self.process_inline_formatting(paragraph, ref_text)
    
    def add_missing_references_report(self):
        """Add a report of missing references at the end of the document."""
        if not self.missing_citations:
            return
        
        # Add section for missing references
        self.doc.add_page_break()
        self.add_formatted_paragraph("Referências Não Encontradas", 'Chapter')
        
        # Add explanation
        explanation = ("As seguintes referências foram citadas no texto mas não foram "
                      "encontradas nos arquivos de bibliografia. Elas aparecem destacadas "
                      "em VERMELHO no documento:")
        self.add_formatted_paragraph(explanation, 'Normal')
        
        # Add list of missing references
        for missing_ref in sorted(self.missing_citations):
            paragraph = self.doc.add_paragraph(style='Normal')
            
            # Add bullet point
            bullet_run = paragraph.add_run("• ")
            bullet_run.font.color.rgb = RGBColor(255, 0, 0)  # Red bullet
            
            # Add missing reference in red
            ref_run = paragraph.add_run(missing_ref)
            ref_run.font.color.rgb = RGBColor(255, 0, 0)  # Red text
            
        print(f"Added missing references report with {len(self.missing_citations)} missing citations")
    
    def convert(self):
        """Main conversion function with comprehensive formatting."""
        print("Starting comprehensive LaTeX to Word conversion...")
        
        # Process bibliography
        self.process_bibliography_files()
        print(f"Loaded {len(self.bibliography)} bibliography entries")
        
        # Add title
        title = "Proposition d'un dispositif numérique scénarisé incluant l'intelligence artificielle pour améliorer l'interaction orale en portugais chez les apprenants francophones"
        self.add_formatted_paragraph(title, 'Title')
        
        # Process chapters
        chapter_paths = [
            'chapters/introduction/main.tex',
            'chapters/literature-review/main.tex',
            'chapters/methodology/main.tex',
            'chapters/results/main.tex',
            'chapters/discussion/main.tex',
            'chapters/conclusion/main.tex',
            'chapters/appendix/main.tex'
        ]
        
        for chapter_path in chapter_paths:
            print(f"Processing {chapter_path}...")
            chapter_content = self.read_file_safely(Path(chapter_path))
            if chapter_content:
                chapter_content = self.process_input_command(chapter_content)
                self.process_latex_structure(chapter_content)
                self.doc.add_page_break()
        
        # Add bibliography
        print(f"Adding bibliography with {len(self.citations)} citations")
        self.add_bibliography()
        
        # Add missing references report if any
        if self.missing_citations:
            self.add_missing_references_report()
        
        # Save document
        output_path = "thesis_final_corrected.docx"
        try:
            self.doc.save(output_path)
            print(f"Successfully converted to: {output_path}")
            print(f"Document statistics:")
            print(f"  - Total paragraphs: {len(self.doc.paragraphs)}")
            print(f"  - Bibliography entries: {len(self.bibliography)}")
            print(f"  - Citations found: {len(self.citations)}")
            print(f"  - Missing citations: {len(self.missing_citations)}")
            print(f"  - Chapters processed: {self.section_counters['chapter']}")
            print(f"  - Sections processed: {self.section_counters['section']}")
            
            if self.missing_citations:
                print(f"\n⚠️  MISSING REFERENCES (highlighted in red):")
                for missing_ref in sorted(self.missing_citations):
                    print(f"     - {missing_ref}")
                print(f"   These references appear in RED in the document.")
            
            return True
        except Exception as e:
            print(f"Error saving document: {e}")
            return False


def main():
    """Main function."""
    if not os.path.exists('thesis.tex'):
        print("Error: thesis.tex not found in current directory")
        print("Please run this script from the root of your LaTeX thesis project")
        return 1
    
    converter = LatexToWordConverter()
    success = converter.convert()
    
    if success:
        print("\nConversion completed successfully!")
        print("Enhanced formatting features:")
        print("  ✓ Hierarchical section numbering")
        print("  ✓ Proper heading styles (H1, H2, H3, H4)")
        print("  ✓ Inline formatting (bold, italic)")
        print("  ✓ Academic citations (Author, Year)")
        print("  ✓ Formatted bibliography with hanging indent")
        print("  ✓ Justified text with proper spacing")
        print("  ✓ Chapter numbering")
        print("  ✓ Definition highlighting")
        print("  ✓ Cleaned LaTeX residuals")
        print("\nNote: Manual review recommended for:")
        print("  - Complex tables and figures")
        print("  - Mathematical formulas")
        print("  - Special formatting requirements")
        return 0
    else:
        print("\nConversion failed. Please check the error messages above.")
        return 1


if __name__ == "__main__":
    sys.exit(main())